"""Generate the Verified Financials product & user guide as a styled .docx.

Standalone doc tool — NOT a project dependency. Run with the venv that has
python-docx installed:

    .venv/bin/python scripts/generate_user_guide.py

Writes "Verified Financials - Product and User Guide.docx" to the repo root.
"""

from __future__ import annotations

import struct
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BRAND = RGBColor(0x8A, 0x1C, 0x1C)        # Red Lion brand red
INK = RGBColor(0x1A, 0x23, 0x32)
MUTED = RGBColor(0x5B, 0x6B, 0x7E)
OK = RGBColor(0x1D, 0x7A, 0x46)
HEADER_FILL = "F2E9E9"                      # light brand tint for table headers

IMG_DIR = Path(__file__).resolve().parents[1] / "docs" / "img"


def png_size(path: Path):
    """(width, height) in pixels from a PNG header, or None if not a PNG."""
    try:
        with path.open("rb") as fh:
            head = fh.read(24)
    except OSError:
        return None
    if head[:8] != b"\x89PNG\r\n\x1a\n":
        return None
    return struct.unpack(">II", head[16:24])


def shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def build() -> Document:
    doc = Document()

    # Base styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for lvl, size in ((1, 17), (2, 13.5), (3, 11.5)):
        st = doc.styles[f"Heading {lvl}"]
        st.font.color.rgb = BRAND
        st.font.size = Pt(size)
        st.font.name = "Calibri"
        st.paragraph_format.space_before = Pt(14 if lvl == 1 else 10)
        st.paragraph_format.space_after = Pt(4)

    def para(text="", *, italic=False, bold=False, color=None, size=None, after=6, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = italic
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(after)
        if align is not None:
            p.alignment = align
        return p

    def rich(parts, *, after=6):
        """parts: list of (text, bold) tuples on one paragraph."""
        p = doc.add_paragraph()
        for text, bold in parts:
            r = p.add_run(text)
            r.bold = bold
        p.paragraph_format.space_after = Pt(after)
        return p

    def bullets(items, *, style="List Bullet"):
        for it in items:
            if isinstance(it, tuple):
                p = doc.add_paragraph(style=style)
                lead, rest = it
                r = p.add_run(lead)
                r.bold = True
                p.add_run(rest)
            else:
                doc.add_paragraph(it, style=style)

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = BRAND
            shade(hdr[i], HEADER_FILL)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                p = cells[i].paragraphs[0]
                # allow **bold** lead before a colon in first column
                p.add_run(str(val))
        if widths:
            for r in t.rows:
                for i, w in enumerate(widths):
                    r.cells[i].width = w
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def figure(name, caption):
        path = IMG_DIR / name
        if not path.exists():
            return  # graceful: doc still builds without the image
        size = png_size(path)
        # Fit within the text column; height-bound tall/portrait crops.
        if size and size[0] / size[1] < (6.3 / 4.6):
            doc.add_picture(str(path), height=Inches(4.6))
        else:
            doc.add_picture(str(path), width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED
        cap.paragraph_format.space_after = Pt(12)

    # ---------------------------------------------------------------- title
    if (IMG_DIR / "red-lion-wordmark.png").exists():
        doc.add_picture(str(IMG_DIR / "red-lion-wordmark.png"), width=Inches(2.6))
        doc.paragraphs[-1].paragraph_format.space_after = Pt(10)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Verified Financials")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = BRAND
    sub = doc.add_paragraph()
    rs = sub.add_run("Product Overview & User Guide")
    rs.font.size = Pt(15)
    rs.font.color.rgb = INK
    para(
        "A Red Lion Advisory tool — verify the data, size the loan, watch the covenant.",
        italic=True, color=MUTED, size=11, after=2,
    )
    para("Version 1.0 · June 2026 · Demonstration build on synthetic data", color=MUTED, size=9, after=14)

    # ---------------------------------------------------------------- 1
    doc.add_heading("1. What Verified Financials is", level=1)
    para(
        "Verified Financials is an asset-based-lending (ABL) analysis tool. It takes a "
        "borrower's financial files and answers the three questions a secured lender cares "
        "about most: can we trust the numbers, how much can the borrower safely borrow against "
        "their collateral, and are they staying within the financial covenant their loan "
        "requires? It turns hours of spreadsheet work into a live, auditable, bank-ready view."
    )
    rich([("In one sentence: ", True),
          ("it ingests the borrower's receivables, inventory, and financial statements, "
           "reconciles them, applies the loan agreement's eligibility rules to produce a "
           "borrowing-base certificate, and monitors the coverage covenant with an early "
           "warning — every figure traceable back to its source file.", False)])
    para(
        "The numbers in this build are synthetic — a fictional ~$180M food distributor, "
        '"McLane-Style Distribution Holdings, LLC," modeled on a McLane-type business — so the '
        "tool can be demonstrated end to end without any real client data.",
        color=MUTED,
    )
    figure("overview.png", "The Overview dashboard — KPI tiles and the collateral-to-availability waterfall.")

    # ---------------------------------------------------------------- 2
    doc.add_heading("2. The problem it solves", level=1)
    para(
        "In asset-based lending, how much a company can borrow is recalculated constantly from "
        "messy, multi-file data — receivables agings, inventory listings, trial balances, and "
        "financial statements that rarely agree with each other on the first pass. Today that "
        "work is manual, slow, error-prone, and hard to audit:"
    )
    bullets([
        ("Data you can't fully trust — ", "the same figure differs across files, and nobody "
         "can quickly say which source is right or where a number came from."),
        ("Slow borrowing-base math — ", "eligibility rules (aged, foreign, intercompany, "
         "concentration limits, reserves) are applied by hand in spreadsheets that break."),
        ("Covenants caught too late — ", "coverage ratios are checked after the quarter closes, "
         "leaving no runway to act before a breach."),
        ("No audit trail — ", "when the bank asks \"where did this come from?\", the answer is a "
         "buried cell reference, if it exists at all."),
    ])
    para(
        "Verified Financials replaces that with one consistent, rules-driven engine: reconcile "
        "once, compute the borrowing base the same way every time, monitor the covenant "
        "continuously, and keep a click-through trail from every headline number to the exact "
        "row of the source file it came from."
    )

    # ---------------------------------------------------------------- 3
    doc.add_heading("3. The three core capabilities", level=1)

    doc.add_heading("3.1  Verification & tie-out", level=2)
    para(
        "Cross-checks figures that should agree across every file the borrower submitted (e.g. "
        "balance-sheet A/R vs. the receivables aging detail; book inventory vs. the trial "
        "balance; this year's revenue across two versions of the financials). Anything that "
        "doesn't reconcile within tolerance is flagged as an exception, with the two conflicting "
        "values, the dollar variance, and a link to each source. In the demo it surfaces three "
        "planted discrepancies — A/R off by $200K, inventory by $500K, and a $1.3M revenue "
        "disagreement between two reports."
    )
    figure("verification.png", "Verification & tie-out — each exception shows both conflicting values, the variance vs. tolerance, and a link to each source file.")

    doc.add_heading("3.2  Borrowing-base engine → bank-ready certificate (the centerpiece)", level=2)
    para(
        "Applies the facility's eligibility rules to convert raw collateral into how much can "
        "actually be borrowed, then renders the formal certificate a lender expects. The logic, "
        "step by step:"
    )
    bullets([
        ("Start from gross collateral — ", "total accounts receivable plus inventory."),
        ("Remove ineligibles — ", "aged receivables, foreign/uninsured accounts, intercompany "
         "balances, single-customer concentration above the cap, and (optionally) the whole "
         "balance of any badly past-due customer (\"cross-aging\")."),
        ("Apply advance rates — ", "lend a set percentage of what remains (e.g. 85% of eligible "
         "A/R, 50% of inventory at cost)."),
        ("Subtract reserves — ", "dilution, rent, and priority-payable holdbacks where the "
         "agreement calls for them."),
        ("Cap and net down — ", "take the lesser of the borrowing base and the revolver "
         "commitment, then subtract what's already drawn and any letter-of-credit exposure to "
         "get excess (net) availability — the room left to draw."),
    ])
    para(
        "In the demo: $50M gross collateral steps down to a $27.33M borrowing base; with $22M "
        "drawn, $5.33M of availability remains. The certificate is downloadable as a clean, "
        "print-ready document and reflects any live what-if edits."
    )
    figure("certificate.png", "The Borrowing Base screen — certificate header, the waterfall, and one-click View / Download.")

    doc.add_heading("3.3  FCCR covenant monitor & early warning", level=2)
    para(
        "Tracks the Fixed-Charge Coverage Ratio (FCCR) — roughly, the cash a business generates "
        "after the costs of running it, divided by its fixed financing obligations — against the "
        "minimum the loan requires. The convention used here:"
    )
    rich([("FCCR = ", True),
          ("(EBITDA − unfinanced capex − cash taxes − distributions) ÷ (cash interest + "
           "scheduled principal), measured trailing-twelve-months.", False)])
    para(
        "The tool compares the ratio to the covenant threshold (1.10x in the demo), reports the "
        "headroom, plots the quarter-over-quarter trend with a forward projection, and raises an "
        "early warning while still compliant but inside a cushion band — so the borrower sees a "
        "potential breach coming with time to act. It also supports a springing covenant (tested "
        "only when availability runs thin) and computes the equity cure needed to restore "
        "compliance after a breach."
    )
    figure("fccr.png", "The FCCR computation, the headroom panel, and the covenant trend & projection.")

    # ---------------------------------------------------------------- 4
    doc.add_heading("4. How to use it", level=1)
    para(
        "The app is a single web dashboard. The left rail navigates between sections; the right "
        "rail is the What-if panel; the top bar holds the scenario selector and the guided tour. "
        "Everything recomputes live as you change inputs."
    )

    doc.add_heading("4.1  Scenarios", level=2)
    para(
        "Switch between two built-in datasets from the top-right selector. Baseline is a healthy "
        "quarter (compliant, but with three reconciliation exceptions and thin covenant "
        "headroom). Stress is a distressed quarter under a stricter agreement (covenant breach, "
        "tighter collateral rules). You can also upload your own data (Section 4.7)."
    )

    doc.add_heading("4.2  The What-if panel (live levers)", level=2)
    para(
        "Edit any term of the agreement and watch every figure recompute instantly — no "
        "spreadsheets to rebuild. Key controls:"
    )
    bullets([
        ("Facility — ", "revolver commitment (the credit ceiling), outstanding drawn (borrowed "
         "today), and the collateral as-of date."),
        ("Accounts receivable — ", "advance rate, concentration cap % and basis, and the "
         "cross-aging taint toggle/threshold."),
        ("Inventory — ", "valuation method (% of cost vs. % of NOLV) and advance rate. Switching "
         "to NOLV auto-snaps the advance rate to its typical default, since NOLV already includes "
         "the liquidation haircut."),
        ("Reserves — ", "dilution, rent, and priority-payable holdbacks."),
        ("FCCR covenant — ", "the covenant threshold and the springing-covenant trigger."),
    ])
    para(
        "A \"What-if active\" badge appears whenever you've changed something; Reset returns to "
        "the as-filed figures.", color=MUTED,
    )
    figure("whatif.png", "The What-if panel — edit any term of the agreement and every figure recomputes live.")

    doc.add_heading("4.3  Provenance drill-down", level=2)
    para(
        "Any underlined number is clickable. Clicking it shows the exact source facts behind it — "
        "the file, the row, the value, and the as-of date — so every figure can be traced to "
        "where it came from. This is what makes the output auditable rather than a black box."
    )

    doc.add_heading("4.4  The borrowing-base certificate", level=2)
    para(
        "On the Borrowing Base screen, View / Print opens the formal certificate in a new tab "
        "(print to PDF from there), and Download saves it as a self-contained file. The "
        "certificate always reflects the current scenario and any what-if edits — so you can "
        "negotiate a rule live and hand over a certificate that matches."
    )

    doc.add_heading("4.5  Covenant trend, Compare, and the dashboard", level=2)
    bullets([
        ("Overview — ", "at-a-glance KPI tiles (availability, borrowing base, FCCR, exceptions) "
         "and the collateral-to-availability waterfall."),
        ("FCCR — ", "the computation, the trend chart, and the forward projection of when the "
         "covenant would be breached if the recent pace continued."),
        ("Compare — ", "place any two scenarios (baseline, stress, or your live what-if) "
         "side by side and see every metric's delta."),
        ("Borrowing Base — ", "the certificate, the waterfall, a \"what moves availability most\" "
         "sensitivity (tornado) chart, and a goal-seek (\"what advance rate frees up $X?\")."),
    ])
    figure("sensitivity.png", "\"What moves availability most\" (sensitivity), goal-seek, and the A/R eligibility detail with rule citations.")
    figure("compare.png", "The Compare view — baseline vs. stress side by side, with every metric's delta.")

    doc.add_heading("4.6  Advisor Briefing & \"Ask the data\" (the AI features)", level=2)
    para(
        "The Advisor Briefing writes a plain-English executive memo from the live figures, and "
        "\"Ask the data\" answers free-text questions about them. These are the only AI-powered "
        "parts of the product — see Section 5. They never compute or change any number; they only "
        "explain the numbers the deterministic engines already produced."
    )

    doc.add_heading("4.7  Upload your own data", level=2)
    para(
        "From Upload Data, download the CSV templates, fill them with your figures, and upload. "
        "The app validates the files (columns present, values well-formed, with row-and-column "
        "error messages) and then runs all three engines on your data using the standard rule "
        "set. Bad inputs are caught at upload with a clear message — never a crash."
    )

    doc.add_heading("4.8  Guided tour", level=2)
    para(
        "The \"Guided tour\" button (and the welcome dialog) launches a narrated walk-through that "
        "drives the app through the full story — verification, borrowing base, the certificate, "
        "covenant health, and a stress-case breach — in about two minutes. Ideal for presenting "
        "to someone seeing the tool for the first time."
    )

    # ---------------------------------------------------------------- 5
    doc.add_heading("5. What is AI vs. what is deterministic", level=1)
    para(
        "This distinction matters for trust, so it is worth being precise. Every financial number "
        "in Verified Financials is produced by deterministic, rules-based code — no AI, no "
        "randomness, fully reproducible. AI is used in exactly one place: to write the "
        "plain-English narrative on the Advisor Briefing screen. The AI never calculates."
    )
    figure("briefing.png", "The Advisor Briefing — an AI-written memo (badged \"AI-generated\") grounded in the computed figures, plus \"Ask the data.\"")
    table(
        ["Capability", "How it works", "AI involved?"],
        [
            ["Verification / tie-out", "Rule-based reconciliation engine", "No — deterministic"],
            ["Borrowing-base & certificate", "Eligibility rules from config", "No — deterministic"],
            ["FCCR covenant & projection", "Formula + linear trend", "No — deterministic"],
            ["What-if, sensitivity, goal-seek", "Recomputed by the engines", "No — deterministic"],
            ["Compare & all charts", "Computed figures", "No — deterministic"],
            ["Upload validation", "Schema + value checks", "No — deterministic"],
            ["Advisor Briefing memo", "LLM writes prose from the figures", "Yes — OpenAI"],
            ["\"Ask the data\" Q&A", "LLM answers from the figures", "Yes — OpenAI"],
        ],
    )
    doc.add_heading("How the AI is kept honest", level=3)
    bullets([
        ("Grounded, not generative-from-scratch — ", "the model is handed a compact JSON of the "
         "already-computed figures and is instructed to use only those numbers and never invent "
         "or estimate. It quotes the engines; it does not do math."),
        ("Isolated — ", "all AI access lives in a single module; swapping or removing the "
         "provider touches nothing else."),
        ("Optional with a safe fallback — ", "with no API key configured, the briefing falls back "
         "to a deterministic, rule-written memo built from the same figures, and Q&A shows a "
         "\"set a key\" message. The product is fully functional with the AI turned off."),
        ("Model — ", "OpenAI Chat Completions (default GPT-4o-mini; configurable)."),
    ])
    para(
        "Bottom line: the math is auditable and reproducible; the AI is a writing assistant on top "
        "of it, constrained to the verified figures.", bold=True,
    )

    # ---------------------------------------------------------------- 6
    doc.add_heading("6. Data handling & privacy", level=1)
    bullets([
        ("Demo data is synthetic — ", "no real company or client data is used in this build."),
        ("What leaves the machine — ", "only when the AI briefing is enabled, and only the small "
         "JSON of computed summary figures is sent to the AI provider. The raw uploaded files "
         "(receivables, inventory, statements) never leave."),
        ("Zero-egress option — ", "turn the AI off (no API key) and nothing is sent anywhere; the "
         "rule-written briefing still works."),
        ("Auditability — ", "every figure retains its provenance back to a source file and row, "
         "so outputs can be defended to a lender or auditor."),
    ])
    para(
        "For a real engagement, sending computed figures to a third-party AI service is a "
        "data-handling decision the firm controls via the on/off key.", color=MUTED,
    )

    # ---------------------------------------------------------------- 7
    doc.add_heading("7. Benefits — the value proposition", level=1)
    para("Verified Financials creates value for everyone in the lending relationship.")

    doc.add_heading("For the advisory firm (the operator)", level=3)
    bullets([
        "Turn days of spreadsheet work into minutes, with a repeatable, defensible method.",
        "Walk a prospect through verification, capacity, and covenant risk in one live session.",
        "Differentiate on transparency: every number is traceable, not a black box.",
        "Scale the same rigor across many clients without scaling headcount.",
    ])
    doc.add_heading("For the lender / bank", level=3)
    bullets([
        "Faster, more confident credit decisions on data that has been reconciled and traced.",
        "A consistent borrowing-base certificate computed the same way every period.",
        "Early covenant warnings that create runway to act before a breach becomes a workout.",
        "A clear audit trail that stands up to internal review and examination.",
    ])
    doc.add_heading("For the borrower / company", level=3)
    bullets([
        "See exactly how much you can draw and which rules are constraining you.",
        "Model \"what if\" — a higher advance rate, a paid-down customer — and see the impact live.",
        "Catch reconciliation errors before they cost you availability or credibility.",
        "Understand covenant headroom early enough to manage it.",
    ])
    doc.add_heading("The headline benefits", level=3)
    bullets([
        ("Speed — ", "live recomputation replaces manual spreadsheet cycles."),
        ("Trust — ", "deterministic math plus full provenance; AI is constrained to verified "
         "figures."),
        ("Foresight — ", "covenant early warning and scenario stress-testing."),
        ("Tangibility — ", "a downloadable, bank-ready certificate as the deliverable."),
    ])

    # ---------------------------------------------------------------- 8
    doc.add_heading("8. What it is — and isn't (current limitations)", level=1)
    bullets([
        ("A demonstration on synthetic data — ", "designed to prove the capabilities, not yet a "
         "production system of record."),
        ("Lightweight upload — ", "fixed-schema CSVs today; no fuzzy column-mapping, Excel, or PDF "
         "ingestion yet."),
        ("Simple covenant projection — ", "a straight-line extrapolation of the recent trajectory, "
         "an early-warning heuristic rather than a forecast of actual future results."),
        ("AI is optional and explanatory — ", "it narrates the figures and is not a source of "
         "truth for any number."),
    ])

    # ---------------------------------------------------------------- 9
    doc.add_heading("9. Glossary", level=1)
    table(
        ["Term", "Plain-English meaning"],
        [
            ["Borrowing base", "How much a lender will lend against collateral after the rules."],
            ["Eligible / ineligible", "Collateral that does (or doesn't) count toward the base."],
            ["Advance rate", "The % of eligible collateral the lender will actually lend."],
            ["Concentration cap", "Limit on how much one customer can count, to avoid over-reliance."],
            ["NOLV", "Net Orderly Liquidation Value — what inventory would fetch if sold off."],
            ["Reserves", "Holdbacks (dilution, rent, priority payables) subtracted from the base."],
            ["Excess availability", "Room left to draw = capped base − drawn − L/C exposure."],
            ["FCCR", "Fixed-Charge Coverage Ratio — cash generated ÷ fixed financing costs."],
            ["Covenant", "A minimum (e.g. FCCR ≥ 1.10x) the borrower must maintain."],
            ["Springing covenant", "A covenant tested only when availability falls below a trigger."],
            ["Equity cure", "Cash the owners inject to restore covenant compliance after a breach."],
            ["Provenance", "The traceable source (file, row, date) behind a figure."],
            ["Tie-out", "Confirming the same figure agrees across different files."],
        ],
    )

    # ---------------------------------------------------------------- 10
    doc.add_heading("10. FAQ", level=1)

    def qa(q, a):
        rich([("Q. ", True), (q, True)], after=2)
        para(a, after=8)

    qa("Does the AI calculate the borrowing base or FCCR?",
       "No. All financial math is deterministic, rules-based code. The AI only writes the "
       "narrative briefing and answers questions about figures the engines already computed.")
    qa("Will the AI make up numbers?",
       "It is instructed to use only the provided computed figures and never invent or estimate. "
       "It quotes the engine output rather than performing calculations.")
    qa("Can we run it without sending data to an AI provider?",
       "Yes. With no API key set, the briefing uses a deterministic rule-written memo and nothing "
       "is sent externally. Every other feature is unaffected.")
    qa("Are the numbers reproducible?",
       "Yes. Given the same inputs and rules, the engines produce identical results every time — "
       "the demo figures are locked by regression tests.")
    qa("Can it run on our own data?",
       "Yes, via the Upload Data screen using the provided CSV templates. The current build uses "
       "a fixed schema; broader ingestion (Excel, fuzzy mapping) is on the roadmap.")
    qa("Where do the rules (advance rates, caps, covenant) come from?",
       "From a configuration that represents the loan agreement — \"the agreement as config.\" A "
       "real client's terms would ship as their own version of that file.")

    para()
    para("Verified Financials — prepared by Red Lion Advisory. Demonstration build on synthetic "
         "data; figures are illustrative.", italic=True, color=MUTED, size=9)

    return doc


def main() -> None:
    out = Path(__file__).resolve().parents[1] / "Verified Financials - Product and User Guide.docx"
    build().save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
